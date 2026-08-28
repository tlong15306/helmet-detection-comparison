"""Build Long's final Word section 2.1.2 from verified experiment artifacts."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUTPUT_PATH = PROJECT_ROOT / "report_drafts" / "Long_2.1.2_Huan_luyen_va_Danh_gia.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
INK = RGBColor(0x1F, 0x29, 0x33)


def set_run_font(run, size: float | None = None, bold: bool | None = None, color: RGBColor | None = None, italic: bool | None = None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table_pr = table._tbl.tblPr
    tbl_w = table_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        table_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = table_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        table_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for grid_col, width in zip(grid.gridCol_lst, widths_dxa):
        grid_col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            cell.width = Inches(width / 1440)
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_cell_text(cell, value: str, *, bold: bool = False, align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT, color: RGBColor = INK, size: float = 9.5) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.05
    run = paragraph.add_run(value)
    set_run_font(run, size=size, bold=bold, color=color)


def add_table(doc: Document, caption: str, headers: list[str], rows: list[list[str]], widths_dxa: list[int]) -> None:
    cap = doc.add_paragraph()
    cap.paragraph_format.space_before = Pt(8)
    cap.paragraph_format.space_after = Pt(4)
    cap.paragraph_format.keep_with_next = True
    cap_run = cap.add_run(caption)
    set_run_font(cap_run, size=10.5, bold=True, color=RGBColor(0, 0, 0))

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths_dxa)
    set_repeat_table_header(table.rows[0])
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_text(cell, header, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=9)
    for row_values in rows:
        row_cells = table.add_row().cells
        for index, (cell, value) in enumerate(zip(row_cells, row_values)):
            align = WD_ALIGN_PARAGRAPH.CENTER if index > 0 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cell, value, align=align, size=9.2)
    for row in table.rows:
        row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))


def add_heading(doc: Document, text: str, level: int) -> None:
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    paragraph.add_run(text)


def add_body(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Inches(0.25)
    paragraph.add_run(text)


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.add_run(text)


def add_equation(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    set_run_font(run, size=11, italic=True, color=RGBColor(0, 0, 0))


def add_figure(doc: Document, image_path: Path, caption: str) -> None:
    figure = doc.add_paragraph()
    figure.alignment = WD_ALIGN_PARAGRAPH.CENTER
    figure.paragraph_format.space_before = Pt(8)
    figure.paragraph_format.space_after = Pt(3)
    figure.paragraph_format.keep_with_next = True
    shape = figure.add_run().add_picture(str(image_path), width=Inches(6.2))
    shape._inline.docPr.set("descr", caption)
    shape._inline.docPr.set("title", image_path.stem)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(7)
    cap_run = cap.add_run(caption)
    set_run_font(cap_run, size=9.5, italic=True, color=RGBColor(0x44, 0x44, 0x44))


def add_figure_pair(doc: Document, left_path: Path, right_path: Path, caption: str) -> None:
    figure = doc.add_paragraph()
    figure.alignment = WD_ALIGN_PARAGRAPH.CENTER
    figure.paragraph_format.space_before = Pt(8)
    figure.paragraph_format.space_after = Pt(3)
    figure.paragraph_format.keep_with_next = True
    for index, image_path in enumerate((left_path, right_path)):
        if index:
            figure.add_run(" ")
        shape = figure.add_run().add_picture(str(image_path), width=Inches(3.02))
        shape._inline.docPr.set("descr", caption)
        shape._inline.docPr.set("title", image_path.stem)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.keep_with_next = False
    cap.paragraph_format.space_after = Pt(7)
    cap_run = cap.add_run(caption)
    set_run_font(cap_run, size=9.5, italic=True, color=RGBColor(0x44, 0x44, 0x44))


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    for level, size, color, before, after in [
        (1, 16, BLUE, 18, 10),
        (2, 13, BLUE, 12, 6),
        (3, 12, DARK_BLUE, 8, 4),
    ]:
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ["List Bullet", "List Number"]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.208

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("BTL TRÍ TUỆ NHÂN TẠO | Faster R-CNN và RetinaNet")
    set_run_font(run, size=8.5, color=RGBColor(0x66, 0x66, 0x66))

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    left = p.add_run("Nguyễn Thành Long | Mục 2.1.2 | Trang ")
    set_run_font(left, size=8.5, color=RGBColor(0x66, 0x66, 0x66))
    add_page_number(p)


def build_document() -> None:
    doc = Document()
    configure_document(doc)
    chart_dir = PROJECT_ROOT / "report_drafts" / "assets" / "png"
    demo_dir = PROJECT_ROOT / "report_drafts" / "assets" / "demo"

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(20)
    title.paragraph_format.space_after = Pt(5)
    run = title.add_run("2.1.2. HUẤN LUYỆN VÀ ĐÁNH GIÁ")
    set_run_font(run, size=18, bold=True, color=RGBColor.from_string(DARK_BLUE))

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(16)
    run = subtitle.add_run("Đề tài: Ứng dụng và so sánh Faster R-CNN và RetinaNet trong phát hiện người điều khiển xe máy không đội mũ bảo hiểm từ hình ảnh giao thông")
    set_run_font(run, size=10.5, italic=True, color=RGBColor(0x55, 0x55, 0x55))

    for label, value in [
        ("Người phụ trách", "Nguyễn Thành Long"),
        ("Phạm vi", "Fine-tune, giao thức đánh giá và diễn giải kết quả thực nghiệm"),
    ]:
        metadata = doc.add_paragraph()
        metadata.paragraph_format.space_after = Pt(2)
        metadata.paragraph_format.line_spacing = 1.05
        label_run = metadata.add_run(f"{label}: ")
        set_run_font(label_run, size=9.5, bold=True, color=RGBColor(0x55, 0x55, 0x55))
        value_run = metadata.add_run(value)
        set_run_font(value_run, size=9.5, color=RGBColor(0x55, 0x55, 0x55))

    add_heading(doc, "2.1.2.1. Mục đích và nguyên tắc thực nghiệm", 2)
    add_body(doc, "Mục này trình bày quy trình tinh chỉnh hai mô hình Faster R-CNN và RetinaNet cho bài toán phát hiện đối tượng trong ảnh giao thông. Mỗi dự đoán gồm khung giới hạn, nhãn lớp và độ tin cậy. Dữ liệu đã xử lý sử dụng ba lớp đối tượng: BikeWithRider, NoHelmet và Helmet; lớp nền chỉ được dùng nội bộ bởi mô hình và không được tính là lớp phát hiện.")
    add_body(doc, "Để bảo đảm phép so sánh công bằng, hai mô hình được khởi tạo từ trọng số COCO do Torchvision cung cấp, dùng cùng bộ dữ liệu, cùng cách chia train/validation/test, cùng phần cứng và cùng giao thức đánh giá. Tập kiểm thử được giữ riêng trong suốt quá trình chọn cấu hình. Checkpoint tốt nhất của mỗi mô hình được chọn theo mAP@0.5:0.95 trên tập xác thực, sau đó chỉ được đánh giá một lần trên tập kiểm thử.")

    add_heading(doc, "2.1.2.2. Quy trình fine-tune", 2)
    add_heading(doc, "a) Chuẩn bị dữ liệu", 3)
    add_body(doc, "Tập dữ liệu được chuyển về định dạng phù hợp với Torchvision và cố định bằng seed 42. Toàn bộ 2.392 ảnh chứa 8.274 khung giới hạn. Cùng một manifest chia tập được dùng cho cả Faster R-CNN và RetinaNet, nhờ đó không có ảnh của train hoặc validation xuất hiện trong test.")
    add_table(
        doc,
        "Bảng 2.1. Phân chia dữ liệu dùng trong thực nghiệm.",
        ["Tập dữ liệu", "Số ảnh", "Số khung giới hạn", "Vai trò"],
        [
            ["Huấn luyện", "1.673", "5.800", "Cập nhật trọng số mô hình"],
            ["Xác thực", "360", "1.236", "Theo dõi huấn luyện và chọn checkpoint"],
            ["Kiểm thử", "359", "1.238", "Đánh giá cuối cùng"],
            ["Tổng", "2.392", "8.274", "-"],
        ],
        [1700, 1100, 1700, 4860],
    )

    add_heading(doc, "b) Cấu hình huấn luyện", 3)
    add_body(doc, "Faster R-CNN là detector hai giai đoạn, còn RetinaNet là detector một giai đoạn sử dụng Focal Loss. Nhóm sử dụng các biến thể có sẵn trong Torchvision và thay đầu dự đoán để phù hợp với ba lớp đối tượng của bài toán. Các siêu tham số chính được giữ nhất quán giữa hai lượt huấn luyện.")
    add_table(
        doc,
        "Bảng 2.2. Cấu hình fine-tune áp dụng cho cả hai mô hình.",
        ["Thành phần", "Giá trị"],
        [
            ["Framework", "PyTorch 2.5.1+cu121; Torchvision 0.20.1+cu121"],
            ["Trọng số khởi tạo", "Trọng số COCO mặc định của Torchvision"],
            ["Số lớp đầu ra", "4, gồm 3 lớp đối tượng và lớp nền nội bộ"],
            ["Optimizer", "SGD; momentum = 0,9; weight decay = 0,0005"],
            ["Learning rate ban đầu", "0,0025"],
            ["Scheduler", "StepLR; gamma = 0,1; step size = 7 epoch"],
            ["Batch size / số epoch", "1 / 20"],
            ["Kích thước ảnh", "Cạnh ngắn 512; cạnh dài tối đa 768 pixel"],
            ["Backbone fine-tune", "Ba tầng cuối của backbone"],
            ["Phần cứng", "Intel Core i5-12450H; RTX 2050 4 GB; RAM 16 GB"],
        ],
        [2900, 6460],
    )
    add_body(doc, "Ở mỗi epoch, mô hình nhận ảnh và annotation của tập huấn luyện, tính loss, lan truyền ngược và cập nhật trọng số qua SGD. Sau epoch, mô hình được suy luận trên tập xác thực; mAP@0.5:0.95 được dùng làm tiêu chí lưu checkpoint tốt nhất. Training loss chỉ phản ánh mục tiêu tối ưu trên train, vì vậy không được dùng thay cho các chỉ số đánh giá tổng quát hóa trên validation hoặc test.")
    add_table(
        doc,
        "Bảng 2.3. Checkpoint được chọn sau fine-tune.",
        ["Mô hình", "Epoch checkpoint tốt nhất", "Thời gian huấn luyện 20 epoch"],
        [
            ["Faster R-CNN", "9", "2 giờ 36 phút 21 giây"],
            ["RetinaNet", "8", "2 giờ 14 phút 44 giây"],
        ],
        [3000, 2900, 3460],
    )

    add_heading(doc, "2.1.2.3. Các chỉ số đánh giá", 2)
    add_heading(doc, "a) Intersection over Union", 3)
    add_body(doc, "Intersection over Union (IoU) đo độ chồng lấp giữa khung dự đoán Bp và khung nhãn thật Bg. IoU có giá trị trong khoảng từ 0 đến 1; giá trị càng cao cho thấy vị trí và kích thước khung dự đoán càng gần nhãn thật.")
    add_equation(doc, "IoU(Bp, Bg) = |Bp ∩ Bg| / |Bp ∪ Bg|")
    add_body(doc, "Khi tính Precision và Recall trong thực nghiệm này, một dự đoán chỉ được ghép với một nhãn thật chưa ghép nếu cùng lớp và có IoU không nhỏ hơn 0,50.")
    add_heading(doc, "b) Precision và Recall", 3)
    add_body(doc, "Sau khi sắp xếp dự đoán theo độ tin cậy giảm dần, evaluator thực hiện ghép một-một theo cách tham lam trong từng lớp. True Positive (TP) là dự đoán đúng lớp và ghép được với một nhãn thật; False Positive (FP) là dự đoán sai lớp, IoU dưới ngưỡng hoặc trùng với nhãn đã ghép; False Negative (FN) là nhãn thật không được ghép với dự đoán nào.")
    add_equation(doc, "Precision = TP / (TP + FP)        Recall = TP / (TP + FN)")
    add_body(doc, "Precision cao cho biết các khung dự đoán có ít báo động sai, trong khi Recall cao cho biết mô hình bỏ sót ít đối tượng. Hai chỉ số thường đánh đổi khi thay đổi ngưỡng confidence; do đó ngưỡng IoU và confidence phải luôn được nêu rõ khi diễn giải kết quả.")
    add_body(doc, "Đối với ứng dụng demo, confidence threshold được lựa chọn trên 360 ảnh validation, không dùng tập test. Nhóm quét các ngưỡng từ 0,05 đến 0,95 với bước 0,05 và chọn theo F1 của lớp NoHelmet; khi bằng F1, lần lượt ưu tiên Recall, Precision và ngưỡng cao hơn. Kết quả chọn 0,85 cho Faster R-CNN (F1 = 0,8631; Precision = 0,8869; Recall = 0,8405) và 0,60 cho RetinaNet (F1 = 0,8216; Precision = 0,8697; Recall = 0,7786). Các ngưỡng này chỉ phục vụ lọc dự đoán khi hiển thị trong demo, không làm thay đổi mAP đã báo cáo.")
    add_heading(doc, "c) AP và mAP", 3)
    add_body(doc, "Average Precision (AP) là diện tích dưới đường Precision-Recall của một lớp tại quy tắc IoU xác định. Mean Average Precision (mAP) là trung bình AP trên các lớp. Theo giao thức COCO, mAP@0.5:0.95 là trung bình AP tại các ngưỡng IoU từ 0,50 đến 0,95 với bước 0,05. Báo cáo đồng thời trình bày mAP@0.5, mAP@0.75 và mAR@100 để quan sát chất lượng phát hiện trong các mức yêu cầu định vị khác nhau.")

    add_heading(doc, "2.1.2.4. Kết quả đánh giá trên tập kiểm thử", 2)
    add_body(doc, "Sau khi chốt checkpoint từ validation, mỗi mô hình được đánh giá một lần trên cùng 359 ảnh test. AP/mAP được tính theo COCO mAP@[IoU=0.50:0.95, bước 0.05] bằng pycocotools. Precision/Recall sử dụng greedy matching, cùng lớp, IoU = 0,50 và không đặt thêm ngưỡng confidence tại evaluator. Hai lần đánh giá dùng cùng tập test, batch size bằng 1 và seed 42.")
    add_table(
        doc,
        "Bảng 2.4. Kết quả phát hiện trên tập kiểm thử; giá trị cao hơn là tốt hơn.",
        ["Mô hình", "mAP@0.5:0.95", "mAP@0.5", "mAP@0.75", "mAR@100"],
        [
            ["Faster R-CNN", "0,6562", "0,9070", "0,7400", "0,7317"],
            ["RetinaNet", "0,6472", "0,8990", "0,7457", "0,7436"],
        ],
        [2400, 1800, 1500, 1600, 2060],
    )
    add_figure(
        doc,
        chart_dir / "test_metrics_comparison.png",
        "Hình 2.1. So sánh các chỉ số phát hiện trên tập kiểm thử. Nguồn: Nhóm tác giả xây dựng từ kết quả đánh giá.",
    )
    add_body(doc, "Faster R-CNN đạt mAP@0.5:0.95 cao hơn RetinaNet 0,0090 điểm, tương đương khoảng 0,9 điểm phần trăm, và mAP@0.5 cao hơn 0,0080 điểm. RetinaNet lại cao hơn tại mAP@0.75 (0,0057 điểm) và mAR@100 (0,0120 điểm). Vì chênh lệch mAP@0.5:0.95 nhỏ và kết quả hiện được quan sát ở một seed, số liệu này chỉ phản ánh xu hướng của cấu hình hiện tại; chưa đủ để khẳng định một mô hình vượt trội hoàn toàn.")
    add_table(
        doc,
        "Bảng 2.5. Phân tích lớp NoHelmet tại IoU = 0,50; không đặt thêm ngưỡng confidence.",
        ["Mô hình", "Precision", "Recall", "AP@0.5:0.95"],
        [
            ["Faster R-CNN", "0,6747", "0,9336", "0,5584"],
            ["RetinaNet", "0,1504", "0,9645", "0,5386"],
        ],
        [3000, 1900, 1900, 2560],
    )
    add_body(doc, "Ở đầu ra thô theo giao thức trên, RetinaNet có Recall của lớp NoHelmet cao hơn nhưng Precision thấp hơn đáng kể. Điều này cho thấy mô hình tạo nhiều dự đoán dương tính sai hơn. Ngưỡng hiển thị trong ứng dụng demo, nếu có, phải được chọn trên tập validation thay vì điều chỉnh theo tập test.")

    add_heading(doc, "2.1.2.5. Đánh giá tốc độ suy luận", 2)
    add_body(doc, "Tốc độ được benchmark trên 100 ảnh validation sau 20 ảnh warm-up, batch size bằng 1 và GPU RTX 2050. Thời gian bao gồm chuyển tensor từ CPU sang GPU, biến đổi nội bộ của Torchvision, forward và hậu xử lý/NMS; không bao gồm đọc ảnh từ ổ đĩa, ghi tệp hoặc vẽ giao diện.")
    add_table(
        doc,
        "Bảng 2.6. Kết quả benchmark suy luận trên tập xác thực.",
        ["Mô hình", "Latency TB (ms/ảnh)", "Median", "P95", "FPS", "Bộ nhớ GPU đỉnh"],
        [
            ["Faster R-CNN", "163,59", "163,80", "175,64", "6,11", "475,9 MiB"],
            ["RetinaNet", "75,24", "75,41", "83,58", "13,29", "376,1 MiB"],
        ],
        [2100, 1640, 1250, 1250, 1100, 2020],
    )
    add_figure(
        doc,
        chart_dir / "latency_fps_comparison.png",
        "Hình 2.2. So sánh latency và FPS suy luận trên tập xác thực. Nguồn: Nhóm tác giả xây dựng từ benchmark.",
    )
    add_body(doc, "Trong giao thức benchmark này, RetinaNet đạt FPS cao hơn khoảng 2,17 lần Faster R-CNN và dùng ít bộ nhớ GPU cấp phát đỉnh hơn. Tuy nhiên, 13,29 FPS chưa được gọi là thời gian thực vì nhóm chưa xác lập tiêu chí FPS/latency cho thuật ngữ đó. Các số đo cũng không đại diện cho toàn bộ ứng dụng demo do đã loại trừ I/O và thời gian vẽ giao diện.")

    add_heading(doc, "2.1.2.6. Phân tích định tính trên ảnh demo", 2)
    add_body(doc, "Trên ảnh một người điều khiển xe máy có đội mũ, cả Faster R-CNN và RetinaNet đều trả đúng hai đối tượng gồm một BikeWithRider và một Helmet tại ngưỡng demo đã chọn bằng validation. Tầng liên kết vai trò v2 ghép duy nhất box đầu với vùng xe và hiển thị “tài xế theo quy tắc · có mũ”; vì vậy không tạo cảnh báo tài xế không đội mũ.")
    add_figure_pair(
        doc,
        demo_dir / "single_rider_faster_rcnn.png",
        demo_dir / "single_rider_retinanet.png",
        "Hình 2.3. Kết quả Faster R-CNN (trái) và RetinaNet (phải) trên ảnh một người điều khiển có đội mũ. Nguồn: Nhóm tác giả chạy ứng dụng demo.",
    )
    add_figure_pair(
        doc,
        demo_dir / "single_rider_faster_role_v2.png",
        demo_dir / "single_rider_retina_role_v2.png",
        "Hình 2.4. Tầng liên kết vai trò v2 của Faster R-CNN (trái) và RetinaNet (phải). Nguồn: Nhóm tác giả chạy ứng dụng demo.",
    )
    add_body(doc, "Độ trễ hiển thị trên hai lần bấm đơn lẻ lần lượt là 588,8 ms với Faster R-CNN và 353,6 ms với RetinaNet. Các giá trị này có thể chịu ảnh hưởng của lần nạp checkpoint và toàn bộ luồng ứng dụng, nên chỉ dùng để mô tả trải nghiệm demo; Bảng 2.6 với 100 ảnh sau warm-up vẫn là căn cứ so sánh tốc độ chính thức.")
    add_body(doc, "Ở cảnh đông người, Faster R-CNN hiển thị 12 detection còn RetinaNet hiển thị 11 detection tại hai ngưỡng validation tương ứng. Không thể kết luận mô hình nào tốt hơn chỉ từ số detection của một ảnh, vì detection bổ sung có thể là phát hiện đúng hoặc false positive. Các box xe và đầu chồng lấn cũng cho thấy giới hạn của dữ liệu ba lớp: khi một vùng xe chứa nhiều box đầu, quy tắc v2 chủ động trả vai trò “chưa xác định” thay vì ép chọn tài xế.")
    add_figure_pair(
        doc,
        demo_dir / "crowded_faster_rcnn.png",
        demo_dir / "crowded_retinanet.png",
        "Hình 2.5. Kết quả Faster R-CNN (trái) và RetinaNet (phải) trên cảnh đông người. Nguồn: Nhóm tác giả chạy ứng dụng demo.",
    )

    add_heading(doc, "2.1.2.7. Nhận xét", 2)
    add_body(doc, "Quy trình thực nghiệm đã giữ cố định dữ liệu và giao thức đánh giá cho hai mô hình. Trên cấu hình hiện tại, Faster R-CNN có mAP@0.5:0.95 nhỉnh hơn nhẹ, trong khi RetinaNet có mAR@100, mAP@0.75 và tốc độ suy luận cao hơn. Lựa chọn mô hình triển khai cần dựa trên mục tiêu sử dụng cụ thể: ưu tiên chất lượng mAP tổng quát, khả năng thu hồi đối tượng, hoặc tốc độ và mức dùng bộ nhớ. Kết luận triển khai cuối cùng cần kết hợp thêm phân tích định tính các ảnh dự đoán đúng/sai và yêu cầu của ứng dụng demo.")

    add_heading(doc, "Tài liệu tham khảo", 2)
    references = [
        "[1] Ren, S., He, K., Girshick, R. và Sun, J. (2015). Faster R-CNN: Towards Real-Time Object Detection with Region Proposal Networks. arXiv:1506.01497.",
        "[2] Lin, T.-Y., Goyal, P., Girshick, R., He, K. và Dollár, P. (2017). Focal Loss for Dense Object Detection. arXiv:1708.02002.",
        "[3] Lin, T.-Y. và cộng sự (2014). Microsoft COCO: Common Objects in Context. arXiv:1405.0312.",
        "[4] Artifact thực nghiệm của nhóm: run_manifest.json, test_metrics.json, test_comparison.json và latency_validation.json trong kho mã nguồn helmet_detection_project.",
    ]
    for reference in references:
        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(4)
        p.add_run(reference)

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.core_properties.title = "2.1.2 Huấn luyện và đánh giá"
    doc.core_properties.author = "Nguyễn Thành Long"
    doc.core_properties.subject = "Bài tập lớn môn Trí tuệ nhân tạo"
    doc.save(OUTPUT_PATH)
    print(f"Created: {OUTPUT_PATH}")


if __name__ == "__main__":
    build_document()
